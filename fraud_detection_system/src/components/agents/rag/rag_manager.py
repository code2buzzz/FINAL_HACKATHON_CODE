import os
import pandas as pd
from docx import Document
from pypdf import PdfReader
from langchain_chroma import Chroma
from langchain_core.documents import Document as LCDocument
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from config.settings import RAG_DATA_ROOT_DIR, ROOT_DIR


class RAG_Manager:
    """
    Unified RAG Service

    Responsibilities:
    - Document ingestion
    - Chunking
    - Embedding generation
    - Chroma persistence
    - Semantic retrieval
    """

    _instance = None
    _initialized = False

    def __new__(cls):

        if cls._instance is None:
            cls._instance = super().__new__(cls)

        return cls._instance

    def __init__(self):

        if RAG_Manager._initialized:
            return

        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
        )

        self.persist_directory = os.path.join(
            ROOT_DIR,
            "storage",
            "chroma_storage",
        )

        self.vector_store = Chroma(
            collection_name="fraud_documents_knowledge_base",
            embedding_function=self.embeddings,
            persist_directory=self.persist_directory,
        )

        RAG_Manager._initialized = True

    # ==================================================
    # DOCUMENT PROCESSING
    # ==================================================

    def process_file_to_documents(
        self,
        file_path: str,
        folder_name: str,
        file_name: str,
    ):
        """
        Converts a file into LangChain documents.
        """

        ext = os.path.splitext(file_path)[1].lower()

        lc_documents = []

        try:

            # ------------------------------------------
            # Structured Data
            # ------------------------------------------

            if ext in [".csv", ".xlsx", ".xls"]:

                df = (
                    pd.read_csv(file_path)
                    if ext == ".csv"
                    else pd.read_excel(file_path)
                )

                for idx, row in df.iterrows():

                    items = [
                        f"{col}: {val}" for col, val in row.items() if pd.notna(val)
                    ]

                    row_string = ", ".join(items)

                    if row_string.strip():

                        lc_documents.append(
                            LCDocument(
                                page_content=row_string,
                                metadata={
                                    "source_file": file_name,
                                    "folder_category": folder_name,
                                    "row_index": idx,
                                    "data_type": "structured_record",
                                },
                            )
                        )

            # ------------------------------------------
            # Unstructured Data
            # ------------------------------------------

            else:

                text_content = ""

                if ext == ".txt":

                    with open(
                        file_path,
                        "r",
                        encoding="utf-8",
                        errors="ignore",
                    ) as f:
                        text_content = f.read()

                elif ext == ".pdf":

                    reader = PdfReader(file_path)

                    text_content = "\n".join(
                        [
                            page.extract_text()
                            for page in reader.pages
                            if page.extract_text()
                        ]
                    )

                elif ext in [".docx", ".doc"]:

                    document = Document(file_path)

                    text_content = "\n".join(
                        paragraph.text for paragraph in document.paragraphs
                    )

                if text_content.strip():

                    chunks = self.text_splitter.split_text(text_content.strip())

                    for idx, chunk in enumerate(chunks):

                        lc_documents.append(
                            LCDocument(
                                page_content=chunk,
                                metadata={
                                    "source_file": file_name,
                                    "folder_category": folder_name,
                                    "chunk_index": idx,
                                    "data_type": "unstructured_text",
                                },
                            )
                        )

        except Exception as e:

            print(f"❌ Error processing file " f"{file_path}: {str(e)}")

        return lc_documents

    # ==================================================
    # INGESTION
    # ==================================================

    def ingest_folders(self, target_folders):

        total_documents = 0

        for folder in target_folders:

            folder_path = os.path.join(
                RAG_DATA_ROOT_DIR,
                folder,
            )

            if not os.path.exists(folder_path):

                print(f"⚠️ Directory not found: " f"{folder_path}")

                continue

            print(f"\n📂 Processing Folder: {folder}")

            files_to_process = [
                file_name
                for file_name in os.listdir(folder_path)
                if os.path.isfile(os.path.join(folder_path, file_name))
                and not file_name.startswith(".")
            ]

            for file_name in files_to_process:

                file_path = os.path.join(
                    folder_path,
                    file_name,
                )

                print(f"   📄 {file_name}")

                documents = self.process_file_to_documents(
                    file_path=file_path,
                    folder_name=folder,
                    file_name=file_name,
                )

                if documents:

                    self.vector_store.add_documents(documents)

                    total_documents += len(documents)

                    print(f"   ✅ Indexed " f"{len(documents)} documents")

        print(
            f"\n🎯 Ingestion Complete | " f"Total Documents Indexed: {total_documents}"
        )

    # ==================================================
    # BASIC RETRIEVAL
    # ==================================================

    def search(
        self,
        query: str,
        category: str,
        k: int = 5,
    ) -> str:

        results = self.vector_store.similarity_search(
            query=query,
            k=k,
            filter={
                "folder_category": category,
            },
        )

        return self._format_results(
            results,
            category,
        )

    # ==================================================
    # RETRIEVAL WITH SCORES
    # ==================================================

    def search_with_score(
        self,
        query: str,
        category: str,
        k: int = 5,
    ):

        results = self.vector_store.similarity_search_with_score(
            query=query,
            k=k,
            filter={
                "folder_category": category,
            },
        )

        return [
            {
                "content": doc.page_content,
                "metadata": doc.metadata,
                "score": float(score),
            }
            for doc, score in results
        ]

    # ==================================================
    # FORMAT RESULTS
    # ==================================================

    def _format_results(
        self,
        docs,
        category: str,
    ):

        if not docs:

            return f"No relevant context found " f"for category '{category}'."

        formatted = []

        for idx, doc in enumerate(docs, start=1):

            meta = doc.metadata

            formatted.append(f"""
                [Evidence {idx}]
                Content: {doc.page_content}
                Source: {meta.get('source_file', 'unknown')}
                Type: {meta.get('data_type', 'unknown')}
                Category: {meta.get('folder_category', 'unknown')}
                """)

        return "\n".join(formatted)
