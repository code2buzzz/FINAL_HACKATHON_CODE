import os
import pandas as pd
from pypdf import PdfReader
from docx import Document
from dotenv import load_dotenv

# Modern standalone imports
from langchain_huggingface import HuggingFaceEmbeddings, HuggingFaceEndpointEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from config.settings import HF_TOKEN
from config.settings import RAG_DATA_ROOT_DIR, ROOT_DIR


class RAG_Manager:
    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(
            model="sentence-transformers/all-MiniLM-L6-v2"
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=200
        )
        self.persist_directory = os.path.join(ROOT_DIR, "storage", "chroma_storage")

    def process_file_to_documents(self, file_path, folder_name, file_name):
        """Parses files and returns a list of individual LangChain Document objects

        based on whether the underlying file is structured (tabular) or unstructured.
        """
        ext = os.path.splitext(file_path)[1].lower()
        lc_documents = []

        try:
            # 📊 STRATEGY 1: TABULAR DATA (Isolate each individual row into its own document)
            if ext in [".csv", ".xlsx", ".xls"]:
                df = (
                    pd.read_csv(file_path)
                    if ext == ".csv"
                    else pd.read_excel(file_path)
                )

                for idx, row in df.iterrows():
                    # Convert the individual row columns into a descriptive text sequence
                    items = [
                        f"{col}: {val}" for col, val in row.items() if pd.notna(val)
                    ]
                    row_string = ", ".join(items)

                    if row_string.strip():
                        doc = LCDocument(
                            page_content=row_string,
                            metadata={
                                "source_file": file_name,
                                "folder_category": folder_name,
                                "row_index": idx,
                                "data_type": "structured_record",
                            },
                        )
                        lc_documents.append(doc)

            # 📄 STRATEGY 2: NARRATIVE TEXT (Requires window chunking via text splitter)
            else:
                text_content = ""
                if ext == ".txt":
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        text_content = f.read()
                elif ext == ".pdf":
                    reader = PdfReader(file_path)
                    text_content = "\n".join(
                        [p.extract_text() for p in reader.pages if p.extract_text()]
                    )
                elif ext in [".docx", ".doc"]:
                    doc = Document(file_path)
                    text_content = "\n".join([p.text for p in doc.paragraphs])

                if text_content.strip():
                    chunks = self.text_splitter.split_text(text_content.strip())
                    for i, chunk in enumerate(chunks):
                        doc = LCDocument(
                            page_content=chunk,
                            metadata={
                                "source_file": file_name,
                                "folder_category": folder_name,
                                "chunk_index": i,
                                "data_type": "unstructured_text",
                            },
                        )
                        lc_documents.append(doc)

        except Exception as e:
            print(f"❌ Error processing {file_path}: {str(e)}")

        return lc_documents

    def ingest_folders(self, target_folders):
        for folder in target_folders:

            folder_path = os.path.join(RAG_DATA_ROOT_DIR, folder)

            if not os.path.exists(folder_path):
                print(f"⚠️ Directory '{folder_path}' not found. Skipping...")
                continue

            print(f"\n📂 Processing folder: [{folder_path}]")

            vector_store = Chroma(
                collection_name=folder,
                embedding_function=self.embeddings,
                persist_directory=self.persist_directory,
            )

            files_to_process = [
                f
                for f in os.listdir(folder_path)
                if os.path.isfile(os.path.join(folder_path, f))
                and not f.startswith(".")
            ]

            for file_name in files_to_process:
                file_path = os.path.join(folder_path, file_name)

                print(f"   📄 Processing: {file_name}...")

                lc_documents = self.process_file_to_documents(
                    file_path,
                    folder,
                    file_name,
                )

                if lc_documents:
                    vector_store.add_documents(lc_documents)
                    print(
                        f"   ✅ Successfully indexed " f"{len(lc_documents)} records."
                    )
